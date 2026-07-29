"""Echter docx-Rendering-Nachweis fuer word-export (Fixrunde PR #488, Issue #446).

Review-Fund (PR #488, Runde 2): "AC1: .docx-Erzeugung/Reparaturfreiheit in Word
ist nicht belegt -- nur ein als 'Test-only, kein Ersatz fuer
document-skills:docx' deklarierter Renderer existiert" und "AC2 (teilweise):
[...] dass Zitate/Literaturverzeichnis tatsaechlich im Zielstil im gerenderten
Dokument erscheinen, ist laut Testdocstring bewusst nicht geprueft."

Das stimmte: die Vorrunde hatte den Renderer als ``_render_reference_docx`` IN
DIESE DATEI gelegt, weil es im Repo keinen gab -- bewiesen war damit nur, dass
der Test ein .docx schreiben kann.

Behoben an der Ursache: ``skills/word-export/scripts/render_docx.py`` erzeugt
das Dokument jetzt als Repo-Code (dieselbe Rolle wie ``render_tex.py`` im
LaTeX-Pfad). Diese Datei ruft ausschliesslich diesen Produktionsrenderer auf --
kein Rendering-Code mehr im Test.

Arbeitsteilung der beiden Suiten:
  - tests/test_issue_446_render_pipeline.py faehrt den kompletten
    dokumentierten Aufrufweg (bash-Bloecke aus commands/word.md) und belegt
    AC2 zeichengenau (Literatureintraege im Zielstil landen unveraendert im
    Dokument, geprueft mit zwei verschiedenen Stilen).
  - diese Datei prueft den Renderer direkt gegen einen echten In-Memory-Vault,
    inkl. Grenzfaellen (HEADING_6, leerer Vault).

Mechanisch nachgewiesen:
  - collect_references()/resolve_cite_markers() laufen wirklich gegen einen
    echten Vault und ihr Ergebnis fliesst durch render_docx() in eine
    tatsaechliche Datei.
  - Kapitelueberschriften landen als echte Word-Formatvorlagen ("Heading 1"
    .. "Heading 6"), nicht als manuelles Fett/Groesse -- inkl. Grenzfall
    HEADING_6.
  - Eine echte Word-native TOC-Feldfunktion (w:fldChar/w:instrText) statt
    statischem Text.
  - \\cite{}-Marker sind im gerenderten Fliesstext zu Klartext aufgeloest,
    kein roher Marker mehr sichtbar.
  - Jedes Paper hat einen erkennbaren Eintrag im gerenderten
    Literaturverzeichnis.
  - Die erzeugte Datei laesst sich verlustfrei wieder oeffnen (python-docx-
    Roundtrip) -- ein "Reparaturhinweis" in echtem Word korreliert praktisch
    immer mit einer kaputten Zip-/XML-Struktur, die auch dieser Roundtrip
    aufdecken wuerde.

Die Stil-FORMATIERUNG der Literatureintraege bleibt bewusst ausserhalb des
Renderers: render_docx.py uebernimmt payload["bibliography"] zeichengenau,
formatiert also nicht selbst (keine zweite Stilregel-Implementierung neben
citation-extraction). Dass genau diese Uebernahme zeichengenau passiert, prueft
tests/test_issue_446_render_pipeline.py mit APA- und Harvard-Eintraegen.

Zusaetzlich, wenn `soffice` (LibreOffice) lokal verfuegbar ist: echte
Kompatibilitaets-Konvertierung nach PDF -- derselbe Nachweisweg, den
document-skills:docx selbst fuer die eigene QA vorschreibt. CI hat kein
LibreOffice installiert, daher dort geskippt -- exakt dasselbe Pattern wie
PDFLATEX_AVAILABLE in tests/test_latex_export.py.
"""

from __future__ import annotations

import json
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest

docx = pytest.importorskip("docx", reason="python-docx nicht installiert (uv sync)")

WORKTREE = Path(__file__).parent.parent
SCRIPTS_DIR = WORKTREE / "skills" / "word-export" / "scripts"
LATEX_SCRIPTS_DIR = WORKTREE / "skills" / "latex-export" / "scripts"
CITATION_REFERENCES_DIR = WORKTREE / "skills" / "citation-extraction" / "references"
FIXTURES_DIR = Path(__file__).parent / "fixtures" / "word_export"

sys.path.insert(0, str(SCRIPTS_DIR))
sys.path.insert(0, str(LATEX_SCRIPTS_DIR))

# LibreOffice ist in der CI-Matrix (ubuntu-latest/macos-latest) nicht
# installiert -- etabliertes Pattern siehe PDFLATEX_AVAILABLE in
# tests/test_latex_export.py: dort nur lokal beweiskraeftig.
SOFFICE_AVAILABLE = shutil.which("soffice") is not None


# ---------------------------------------------------------------------------
# Huelle um den PRODUKTIONSRENDERER (kein Rendering-Code im Test)
# ---------------------------------------------------------------------------


def _render_reference_docx(
    chapter_heading: str,
    body_text: str,
    papers: list[dict],
    out_path: Path,
) -> None:
    """Baut die Payload-Form auf und uebergibt sie an render_docx.py.

    `papers` -> minimale Literatureintraege, wie sie die Stilstufe des Agenten
    (Schritt 4 in commands/word.md) in die Payload zurueckschreibt. Die
    Formatierung ist hier absichtlich trivial: geprueft wird, dass der Renderer
    sie uebernimmt, nicht wie sie aussieht (Stilregeln gehoeren zu
    citation-extraction, nicht in diesen Test).
    """
    from render_docx import render_docx

    bibliography = []
    for paper in papers:
        csl = json.loads(paper.get("csl_json", "{}"))
        title = csl.get("title", paper.get("paper_id", "?"))
        authors = csl.get("author", [])
        family = authors[0].get("family") if authors else paper.get("paper_id", "?")
        bibliography.append(f"{family}: {title}")

    render_docx(
        {
            "chapters": [
                {
                    "source": "1-kapitel.md",
                    "path": "kapitel/1-kapitel.md",
                    "body": (
                        f"# {chapter_heading}\n\n{body_text}\n\n"
                        "###### Unterkapitel-Ebene\n\nGrenzfall HEADING_6.\n"
                    ),
                }
            ],
            "papers": papers,
            "bibliography": bibliography,
            "style_file": "apa.md",
            "context": {"Thema": "Testarbeit"},
        },
        out_path,
    )


# ---------------------------------------------------------------------------
# AC1 + AC2 (teilweise) — volle Pipeline einmal wirklich ausgefuehrt
# ---------------------------------------------------------------------------


class TestRealDocxRenderFromPipeline:
    def test_pipeline_output_renders_into_valid_reopenable_docx(self, tmp_path):
        """collect_references()/resolve_cite_markers() -> echtes, wieder
        oeffenbares .docx mit realen Formatvorlagen und aufgeloesten Zitaten."""
        from academic_vault.db import VaultDB
        from academic_vault.server import add_paper
        from collect_references import collect_references, resolve_cite_markers

        db_path = str(tmp_path / "vault.db")
        db = VaultDB(db_path)
        db.init_schema()
        add_paper(
            db_path=db_path,
            paper_id="smith2023",
            csl_json=json.dumps(
                {
                    "title": "DevOps Governance in KMU",
                    "type": "article-journal",
                    "author": [{"family": "Smith", "given": "John"}],
                    "issued": {"date-parts": [[2023]]},
                }
            ),
        )
        add_paper(
            db_path=db_path,
            paper_id="jones2022",
            csl_json=json.dumps(
                {
                    "title": "Cloud-Transformation und Governance",
                    "type": "article-journal",
                    "author": [
                        {"family": "Jones", "given": "Anna"},
                        {"family": "Lee", "given": "Kim"},
                    ],
                    "issued": {"date-parts": [[2022]]},
                }
            ),
        )

        academic_context_text = (FIXTURES_DIR / "academic_context_default.md").read_text()
        result = collect_references(db_path, academic_context_text, CITATION_REFERENCES_DIR)
        # Vorbedingung: die Pipeline lief wirklich gegen einen befuellten Vault.
        assert len(result["papers"]) == 2

        chapter_text = (FIXTURES_DIR / "kapitel_with_cite.md").read_text(encoding="utf-8")
        lines = chapter_text.splitlines()
        chapter_heading = lines[0].lstrip("#").strip()
        chapter_body = "\n".join(lines[1:]).strip()
        resolved_body = resolve_cite_markers(chapter_body, result["papers"])
        assert "\\cite" not in resolved_body  # Vorbedingung fuer den Rendering-Schritt

        out_path = tmp_path / "export.docx"
        _render_reference_docx(chapter_heading, resolved_body, result["papers"], out_path)

        # -- Repair-Proxy: die Datei laesst sich verlustfrei wieder oeffnen --
        reopened = docx.Document(str(out_path))

        headings = {
            p.text: p.style.name for p in reopened.paragraphs if p.style.name.startswith("Heading")
        }
        assert headings.get("Inhaltsverzeichnis") == "Heading 1"
        assert headings.get(chapter_heading) == "Heading 1"
        assert headings.get("Unterkapitel-Ebene") == "Heading 6"
        assert headings.get("Literaturverzeichnis") == "Heading 1"

        body_texts = [p.text for p in reopened.paragraphs]
        assert any("(Smith 2023)" in t for t in body_texts), (
            "Aufgeloester Kurzverweis fehlt im Fliesstext"
        )
        assert any("(Jones & Lee 2022)" in t for t in body_texts)
        assert not any("\\cite" in t for t in body_texts)

        joined = "\n".join(body_texts)
        assert "DevOps Governance in KMU" in joined, (
            "Paper 1 fehlt im gerenderten Literaturverzeichnis"
        )
        assert "Cloud-Transformation und Governance" in joined, (
            "Paper 2 fehlt im gerenderten Literaturverzeichnis"
        )

        # -- echte Word-native TOC-Feldfunktion, kein statischer Text --
        with zipfile.ZipFile(out_path) as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
        assert "<w:instrText" in document_xml and "TOC" in document_xml

    def test_empty_vault_still_renders_valid_docx_without_bibliography_entries(self, tmp_path):
        """Fehlerpfad "Vault leer" (SKILL.md) fuehrt trotzdem zu einem gueltigen,
        wieder oeffenbaren Dokument -- kein Absturz, kein kaputtes Literatur-
        verzeichnis."""
        from academic_vault.db import VaultDB
        from collect_references import collect_references

        db_path = str(tmp_path / "empty_vault.db")
        db = VaultDB(db_path)
        db.init_schema()

        academic_context_text = (FIXTURES_DIR / "academic_context_default.md").read_text()
        result = collect_references(db_path, academic_context_text, CITATION_REFERENCES_DIR)
        assert result["papers"] == []

        out_path = tmp_path / "export_empty.docx"
        _render_reference_docx("Einleitung", "Text ohne Zitate.", result["papers"], out_path)

        reopened = docx.Document(str(out_path))
        assert reopened.paragraphs  # oeffnet ohne Exception, Inhalt vorhanden


# ---------------------------------------------------------------------------
# AC6 — Review-Fund (PR #488, flowkit Runde 2): `add_heading()` ohne
# KeyError-Fallback -- eine per --template geladene Fremdvorlage ohne
# "Title"/"Heading N"-Formatvorlagen liess den Export mit rohem Traceback
# sterben statt der in SKILL.md dokumentierten "FEHLER:"-Meldung.
# ---------------------------------------------------------------------------


class TestTemplateWithoutHeadingStyles:
    def _build_template_without_heading_styles(self, path: Path) -> None:
        """Baut ein gueltiges .docx OHNE 'Title'/'Heading N'-Formatvorlagen --
        simuliert eine Fremdvorlage, die diese Word-Standardstile nicht kennt."""
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE

        template = Document()
        for style in list(template.styles):
            if style.type == WD_STYLE_TYPE.PARAGRAPH and (
                style.name == "Title" or style.name.startswith("Heading")
            ):
                style.element.getparent().remove(style.element)
        template.save(str(path))

    def test_render_docx_falls_back_instead_of_raising_keyerror(self, tmp_path, monkeypatch):
        monkeypatch.setattr(Path, "home", lambda: tmp_path)
        profiles_dir = tmp_path / ".academic-research" / "library-profiles"
        profiles_dir.mkdir(parents=True)
        self._build_template_without_heading_styles(profiles_dir / "leibniz.docx")

        from render_docx import render_docx

        out_path = tmp_path / "export.docx"
        # Vorbedingung: die installierte python-docx-Version wirft hier
        # tatsaechlich KeyError -- sonst waere dieser Test kein Nachweis.
        with pytest.raises(KeyError):
            docx.Document(str(profiles_dir / "leibniz.docx")).add_heading("x", level=0)

        render_docx(
            {
                "chapters": [
                    {"source": "1.md", "path": "kapitel/1.md", "body": "# Kapitel\n\nText.\n"}
                ],
                "papers": [],
                "bibliography": [],
                "style_file": "apa.md",
                "context": {"Thema": "Testarbeit"},
            },
            out_path,
            template="leibniz",
        )

        reopened = docx.Document(str(out_path))
        assert reopened.paragraphs  # kein Absturz -- Fallback griff statt Traceback
        assert any(p.text == "Kapitel" for p in reopened.paragraphs)
        assert any(p.text == "Testarbeit" for p in reopened.paragraphs)


class TestRealDocxRenderMatchesOfficeQAMethod:
    """Staerkerer Repair-Hinweis-Proxy ueber dieselbe QA-Methode, die
    document-skills:docx selbst fuer die eigene Verifikation vorschreibt.
    Nur lokal beweiskraeftig -- siehe Moduldocstring."""

    @pytest.mark.skipif(
        not SOFFICE_AVAILABLE,
        reason="soffice/LibreOffice nicht in PATH (in CI-Matrix nicht installiert)",
    )
    def test_rendered_docx_converts_cleanly_via_soffice(self, tmp_path):
        out_path = tmp_path / "export.docx"
        _render_reference_docx("Einleitung", "Testtext ohne Zitate.", [], out_path)

        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_path),
                str(out_path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        assert result.returncode == 0, (
            f"soffice-Konvertierung fehlgeschlagen (Reparaturhinweis-Signal):\n"
            f"{result.stdout}\n{result.stderr}"
        )
        pdf_path = tmp_path / "export.pdf"
        assert pdf_path.exists() and pdf_path.stat().st_size > 0
