"""render_docx.py — Payload -> echtes .docx (Issue #446, Fixrunde PR #488).

Warum es dieses Skript gibt
---------------------------
Bis zur zweiten Fixrunde endete der word-export-Pfad bei einer JSON-Payload;
das eigentliche Rendern war eine Prosa-Anweisung an den Agenten ("rufe
document-skills:docx auf und rendere daraus"). Damit gab es kein Artefakt, das
die Testsuite ausfuehren konnte -- AC1 (".docx laesst sich in Word ohne
Reparaturhinweis oeffnen") und AC2 ("Zitate/Literaturverzeichnis erscheinen im
Zielstil") waren strukturell unbelegbar, und der Nachweisversuch landete als
test-eigener Renderer in tests/test_word_export_docx_render.py.

Dieses Skript uebernimmt exakt die Rolle, die ``render_tex.py`` im
Geschwister-Skill ``latex-export`` hat: Repo-Code erzeugt die Zieldatei
deterministisch, und genau deshalb ist der Pfad testbar. ``document-skills:docx``
bleibt das deklarierte Word-Backend des Plugins (Praeflight in SKILL.md) und
uebernimmt Layout-Verfeinerung *auf der hier erzeugten Datei* -- es ersetzt sie
nicht mehr.

Was hier bewusst NICHT passiert
-------------------------------
Zitierstil-Formatierung. ``payload["bibliography"]`` enthaelt die vom Agenten
aus ``payload["style_rules"]`` (unveraendert aus ``citation-extraction/
references/<style>.md``) formatierten Eintraege; dieses Skript uebernimmt sie
zeichengenau und in unveraenderter Reihenfolge. Eine zweite Stilregel-
Implementierung neben ``citation-extraction`` entsteht damit nicht. Fehlen die
Eintraege, obwohl der Vault Papers liefert, bricht das Skript ab, statt ein
Literaturverzeichnis in einem unbelegten Format zu schreiben (Preamble-Block
"Keine Fabrikation").

Oeffentliche API:
  render_docx(payload, out_path, template=None) -> list[str]   (Meldungen)
  convert_to_pdf(docx_path) -> tuple[Path | None, list[str]]
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

#: Fehlt python-docx, ist das kein Programmierfehler, sondern eine
#: unvollstaendige Installation -- verstaendliche Meldung statt Traceback (AC6).
_MISSING_DOCX_MESSAGE = (
    "FEHLER: Das Python-Paket 'python-docx' ist nicht installiert - "
    "es wird deshalb keine Word-Datei erzeugt. Nachinstallieren mit "
    "`pip install python-docx` (bzw. `bash scripts/setup.sh` fuer die "
    "komplette Plugin-Umgebung)."
)

__all__ = ["render_docx", "convert_to_pdf", "RenderError", "main"]


class RenderError(RuntimeError):
    """Abbruchgrund, der als 'FEHLER: ...' ausgegeben wird (kein Traceback)."""


# ---------------------------------------------------------------------------
# Markdown-Teilmenge -> Word-Absaetze
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
#: **fett** / *kursiv* / `code` -- alles andere bleibt Klartext.
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)")

#: Word-Formatvorlagen fuer Listen. Ein Fehlen (fremde Vorlage per --template)
#: darf den Export nicht kippen -> Fallback auf den Standard-Absatz.
_BULLET_STYLE = "List Bullet"
_ORDERED_STYLE = "List Number"


def _add_inline_runs(paragraph, text: str) -> None:
    """Schreibt `text` als Runs, mit **fett**/*kursiv* als echte Zeichenformate.

    Ohne diese Umsetzung landeten die Markdown-Sternchen roh im Word-Dokument.
    Ueberschriften bleiben davon unberuehrt -- deren Auszeichnung kommt aus der
    Absatz-Formatvorlage, nicht aus manuellem Fett (AC1).
    """
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            paragraph.add_run(part[1:-1])
        else:
            paragraph.add_run(part)


def _add_styled_paragraph(document, text: str, style: str):
    """Absatz mit `style`; faellt auf den Standard zurueck, wenn die Vorlage fehlt."""
    try:
        paragraph = document.add_paragraph(style=style)
    except KeyError:
        paragraph = document.add_paragraph()
    _add_inline_runs(paragraph, text)
    return paragraph


def _add_heading(document, text: str, level: int):
    """Ueberschrift mit Formatvorlagen-Fallback (Review-Fund PR #488, Runde 2).

    python-docx' `add_heading()` wirft `KeyError`, wenn die zugrunde liegende
    Formatvorlage ("Title" bei level=0, "Heading 1".."Heading 6" sonst) in
    einer per `--template` geladenen Fremdvorlage fehlt -- ohne Abfangen ein
    roher Traceback statt der in SKILL.md dokumentierten `FEHLER:`-Meldung
    (AC6). Fallback analog zu `_add_styled_paragraph`: fett hervorgehobener
    Standardabsatz statt Abbruch. Damit bleibt die Ueberschrift optisch
    erkennbar, auch ohne aktualisiertes Word-Inhaltsverzeichnis fuer diesen
    einen Eintrag.
    """
    try:
        return document.add_heading(text, level=level)
    except KeyError:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        return paragraph


def _add_markdown_body(document, markdown: str) -> None:
    """Rendert die Markdown-Teilmenge der Kapitel in echte Word-Absaetze.

    Ueberschriftenebenen `#`..`######` werden zu den Word-Formatvorlagen
    "Heading 1".."Heading 6" (python-docx `add_heading(level=n)`), NICHT zu
    manuell fett/vergroessertem Text -- nur so kann Word ein automatisches
    Inhaltsverzeichnis aufbauen (AC1).
    """
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            _add_inline_runs(document.add_paragraph(), " ".join(buffer))
            buffer.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush()
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            flush()
            _add_heading(document, heading.group(2).strip(), level=len(heading.group(1)))
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            flush()
            _add_styled_paragraph(document, bullet.group(1).strip(), _BULLET_STYLE)
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            flush()
            _add_styled_paragraph(document, ordered.group(1).strip(), _ORDERED_STYLE)
            continue

        buffer.append(stripped)

    flush()


# ---------------------------------------------------------------------------
# Word-native Feldfunktionen und Dokumentteile
# ---------------------------------------------------------------------------


def _add_toc_field(document) -> None:
    """Echte Word-Feldfunktion fuer das Inhaltsverzeichnis (kein statischer Text).

    Statischer Verzeichnistext liesse sich in Word nicht aktualisieren und
    haette keine Verbindung zu den Formatvorlagen -- AC1 verlangt genau diese
    Verbindung.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = document.add_paragraph().add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-6" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, end):
        run._r.append(element)


#: Sichtbare Platzhalter statt erfundener Angaben: was academic_context.md nicht
#: hergibt, wird als ausfuellbare Leerstelle gerendert (Preamble "Keine
#: Fabrikation"), nicht plausibel geraten.
_TITLE_PAGE_FIELDS = [
    ("Typ", "Art der Arbeit"),
    ("Universität", "Hochschule"),
    ("Studiengang", "Studiengang"),
    ("Betreuer", "Betreuung"),
    ("Abgabetermin", "Abgabetermin"),
]


def _add_title_page(document, context: dict) -> None:
    _add_heading(document, context.get("Thema") or "[Titel der Arbeit]", level=0)
    for key, label in _TITLE_PAGE_FIELDS:
        value = context.get(key)
        document.add_paragraph(f"{label}: {value}" if value else f"{label}: [bitte ergaenzen]")
    document.add_paragraph("Verfasst von: [Name]")
    document.add_page_break()


_DECLARATION_TITLE = "Eidesstattliche Erklärung"
#: Generischer Wortlaut. Hochschulspezifische Formulierungen sind laut Issue
#: #446 ("Out") ein eigenes Issue -- hier wird keiner erfunden.
_DECLARATION_BODY = (
    "Ich versichere hiermit, dass ich die vorliegende Arbeit selbstständig "
    "verfasst und keine anderen als die angegebenen Quellen und Hilfsmittel "
    "benutzt habe. Alle Stellen, die wörtlich oder sinngemäß aus "
    "veröffentlichten oder nicht veröffentlichten Schriften entnommen wurden, "
    "sind als solche kenntlich gemacht."
)


def _add_declaration(document) -> None:
    document.add_page_break()
    _add_heading(document, _DECLARATION_TITLE, level=1)
    document.add_paragraph(_DECLARATION_BODY)
    document.add_paragraph("")
    document.add_paragraph("Ort, Datum")
    document.add_paragraph("Unterschrift")


def _base_document(template: str | None, messages: list[str]):
    """Basisdokument -- optional aus einer Hochschul-Wordvorlage.

    `--template <uni>` sucht `~/.academic-research/library-profiles/<uni>.docx`
    (gemeinsamer Profil-Slot mit `latex-export`, dort die `.tex`-Variante).
    Fehlt die Datei, wird generisch weitergebaut statt abgebrochen -- der in
    SKILL.md dokumentierte Fehlerpfad "Template nicht gefunden".
    """
    import docx

    if not template:
        return docx.Document()

    path = Path.home() / ".academic-research" / "library-profiles" / f"{template}.docx"
    if not path.is_file():
        messages.append(f"Template `{template}` fehlt ({path}) - generisches Titelblatt.")
        return docx.Document()
    return docx.Document(str(path))


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------


def render_docx(payload: dict, out_path: Path | str, template: str | None = None) -> list[str]:
    """Rendert die Payload aus `collect_references.py` in eine echte `.docx`.

    Gibt Hinweismeldungen zurueck (leer = nichts zu melden). Wirft RenderError
    bei Abbruchgruenden, die als "FEHLER: ..." ausgegeben werden.
    """
    papers = payload.get("papers") or []
    bibliography = payload.get("bibliography")
    if papers and not bibliography:
        raise RenderError(
            "Der Vault liefert "
            f"{len(papers)} Literatureintraege, aber die Payload enthaelt kein Feld "
            "'bibliography'. Die Eintraege muessen vor dem Rendern mit den "
            f"Regeln aus '{payload.get('style_file', 'apa.md')}' "
            "(payload['style_rules'], Quelle citation-extraction/references/) "
            "formatiert und als Liste unter 'bibliography' in die Payload "
            "geschrieben werden. Dieses Skript formatiert bewusst nicht selbst - "
            "sonst entstuende eine zweite Stilregel-Implementierung."
        )

    messages: list[str] = []
    document = _base_document(template, messages)

    _add_title_page(document, payload.get("context") or {})

    _add_heading(document, "Inhaltsverzeichnis", level=1)
    _add_toc_field(document)
    document.add_page_break()

    for chapter in payload.get("chapters") or []:
        _add_markdown_body(document, chapter.get("body", ""))

    _add_heading(document, "Literaturverzeichnis", level=1)
    for entry in bibliography or []:
        # Zeichengenau uebernehmen: die Stilentscheidung ist bereits gefallen.
        document.add_paragraph(str(entry))

    _add_declaration(document)

    out = Path(out_path)
    if str(out.parent):
        out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))
    return messages


def convert_to_pdf(docx_path: Path | str) -> tuple[Path | None, list[str]]:
    """Konvertiert die erzeugte `.docx` per LibreOffice -- kein eigener Renderer.

    Fehlt `soffice`, bleibt die `.docx` das gueltige Ergebnis (dokumentierter
    Fehlerpfad, kein Abbruch).
    """
    source = Path(docx_path)
    if shutil.which("soffice") is None:
        return None, [
            "LibreOffice (`soffice`) nicht gefunden - PDF-Konvertierung "
            "uebersprungen, `.docx` verfuegbar."
        ]

    result = subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(source.parent),
            str(source),
        ],
        capture_output=True,
        text=True,
    )
    pdf_path = source.with_suffix(".pdf")
    if result.returncode != 0 or not pdf_path.is_file():
        return None, [
            "PDF-Konvertierung fehlgeschlagen - `.docx` verfuegbar. "
            f"soffice: {(result.stderr or result.stdout).strip()}"
        ]
    return pdf_path, []


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description=(
            "Rendert die Payload aus collect_references.py in eine echte .docx "
            "(Schritt 4 von /academic-research:word, Issue #446)."
        ),
    )
    parser.add_argument("--payload", required=True, help="JSON-Payload aus collect_references.py")
    parser.add_argument("--output", required=True, help="Zieldatei (.docx)")
    parser.add_argument("--format", default="docx", choices=["docx", "pdf"])
    parser.add_argument("--template", default=None, help="Uni-Kuerzel der Hochschul-Wordvorlage")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _build_arg_parser().parse_args(argv)

    try:
        import docx  # noqa: F401
    except ImportError:
        print(_MISSING_DOCX_MESSAGE, file=sys.stderr)
        return 1

    try:
        payload = json.loads(Path(args.payload).read_text(encoding="utf-8"))
    except OSError as exc:
        print(f"FEHLER: Payload nicht lesbar ({exc}).", file=sys.stderr)
        return 1
    except json.JSONDecodeError as exc:
        print(f"FEHLER: Payload '{args.payload}' ist kein gueltiges JSON ({exc}).", file=sys.stderr)
        return 1

    try:
        messages = render_docx(payload, args.output, template=args.template or None)
    except RenderError as exc:
        print(f"FEHLER: {exc}", file=sys.stderr)
        return 1
    except OSError as exc:
        print(f"FEHLER: {exc}", file=sys.stderr)
        return 1

    outputs = [str(args.output)]
    if args.format == "pdf":
        pdf_path, pdf_messages = convert_to_pdf(args.output)
        messages.extend(pdf_messages)
        if pdf_path is not None:
            outputs.append(str(pdf_path))

    chapters = len(payload.get("chapters") or [])
    entries = len(payload.get("bibliography") or [])
    print(f"Erzeugt: {', '.join(outputs)} ({chapters} Kapitel, {entries} Literatureintraege)")
    for message in messages:
        print(message, file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
